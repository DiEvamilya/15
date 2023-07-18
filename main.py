# а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
# б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран.
# в) Создайте абзац с текстом и запишите это в новый word-файл, измените шрифт и размер
# шрифта.

import docx
doc = docx.Document('hello.docx')

text = []

for paragraph in doc.paragraphs:
      print(paragraph.text)

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            print(run.text)

from docx import Document

doc = Document()
doc.add_paragraph('Наш мир ослепительно прекрасен и многообразен.')
doc.save('World.docx')
for pargraph in doc.paragraphs:
    print(pargraph.text)

import docx

doc = docx.Document('World.docx')
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.bold = True
        run.italic = True

doc.save('World1.docx')


